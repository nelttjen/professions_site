import datetime
import time
import xlsxwriter
import re
import sys
import os
import comtypes.client
import requests
import docx
from docx.shared import Inches

from django.shortcuts import render
from django.http import HttpResponseNotFound, HttpResponseBadRequest, FileResponse
from django.utils import timezone

from kwork_order.settings import DOMAIN, MEDIA_ROOT, MEDIA_URL, BASE_DIR
from .models import *

wdFormatPDF = 17
ua = {"User-Agent": 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) '
                    'Chrome/106.0.0.0 YaBrowser/22.11.5.715 Yowser/2.5 Safari/537.36'}
NO_IMAGE_LINK = 'media/default/no_image.png'
HH_ENDPOINT = 'https://api.hh.ru/vacancies/'
last_fetch = timezone.now() - datetime.timedelta(days=333)
cache = []


# Create your views here.
def index(request):
	return render(request, 'app/index.html')


def profs(request):
	profession_list = Profession.objects.filter(is_visible=True).all().values_list('id', 'name', named=True)
	payload = {
		'profs': profession_list
	}
	return render(request, 'app/professions.html', payload)


def prof_overview(request, prof_id):
	if not (profession := Profession.objects.filter(id=prof_id, is_visible=True).first()):
		return HttpResponseNotFound()
	prof_descriptions = Profession.objects.filter(id=prof_id).prefetch_related('professiondescriptionblock_set') \
		.values_list('professiondescriptionblock__title', 'professiondescriptionblock__description', named=True)
	payload = {
		'prof': profession,
		'prof_descs': prof_descriptions
	}
	return render(request, 'app/profession_overview.html', payload)


def prof_view_vostr(request, prof_id):
	if not (profession := Profession.objects.filter(id=prof_id, is_visible=True).first()):
		return HttpResponseNotFound()
	year_diagram = YearDiagram.objects.filter(profession=profession).first()
	vacancies_diagram = VacanciesDiagram.objects.filter(profession=profession).first()
	if year_diagram:
		link = f'media/year_diagram/{year_diagram.id}.jpg'
		try:
			with open(link, 'wb') as file:
				file.write(year_diagram.image)
		except:
			link = NO_IMAGE_LINK
		year_diagram = {
			'title': year_diagram.diagram_title,
			'image': '/' + link
		}
	else:
		year_diagram = {
			'title': 'Диаграмма зарплат по годам не добавлена',
			'image': '/' + NO_IMAGE_LINK
		}
	if vacancies_diagram:
		link = f'media/vacancies_diagram/{vacancies_diagram.id}.jpg'
		try:
			with open(link, 'wb') as file:
				file.write(vacancies_diagram.image)
		except:
			link = NO_IMAGE_LINK
		vacancies_diagram = {
			'title': vacancies_diagram.diagram_title,
			'image': '/' + link
		}
	else:
		vacancies_diagram = {
			'title': 'Диаграмма количества вакансий по годам не добавлена',
			'image': '/' + NO_IMAGE_LINK
		}
	payload = {
		'prof': profession,
		'diag_1': year_diagram,
		'diag_2': vacancies_diagram,
	}
	return render(request, 'app/profession_demand.html', payload)


def prof_view_geo(request, prof_id):
	if not (profession := Profession.objects.filter(id=prof_id, is_visible=True).first()):
		return HttpResponseNotFound()
	cities_diagram = CitiesDiagram.objects.filter(profession=profession).first()
	cities_vacancies_diagram = CitiesVacanciesDiagram.objects.filter(profession=profession).first()
	if cities_diagram:
		link = f'media/year_diagram/{cities_diagram.id}.jpg'
		try:
			with open(link, 'wb') as file:
				file.write(cities_diagram.image)
		except:
			link = NO_IMAGE_LINK
		cities_diagram = {
			'title': cities_diagram.diagram_title,
			'image': '/' + link
		}
	else:
		cities_diagram = {
			'title': 'Диаграмма зарплат по городам не добавлена',
			'image': '/' + NO_IMAGE_LINK
		}
	if cities_vacancies_diagram:
		link = f'media/vacancies_diagram/{cities_vacancies_diagram.id}.jpg'
		try:
			with open(link, 'wb') as file:
				file.write(cities_vacancies_diagram.image)
		except:
			link = NO_IMAGE_LINK
		cities_vacancies_diagram = {
			'title': cities_vacancies_diagram.diagram_title,
			'image': '/' + link
		}
	else:
		cities_vacancies_diagram = {
			'title': 'Диаграмма количества вакансий по городам не добавлена',
			'image': '/' + NO_IMAGE_LINK
		}
	payload = {
		'prof': profession,
		'diag_1': cities_diagram,
		'diag_2': cities_vacancies_diagram,
	}
	return render(request, 'app/profession_geo.html', payload)


def years_view(request, prof_id):
	if not (prof := Profession.objects.filter(id=prof_id, is_visible=True).first()):
		return HttpResponseNotFound()
	years = Year.objects.filter(is_visible=True, profession_id=prof_id).all().order_by('-year'). \
		values_list('year', named=True)
	return render(request, 'app/profession_years.html', {'years': years, 'prof': prof})


def skills_year_view(request, prof_id, year):
	if not (year := Year.objects.filter(profession_id=prof_id, year=year, is_visible=True).first()):
		return HttpResponseNotFound()
	year_skills = Skill.objects.filter(year=year).order_by('-weight').values_list('name', 'weight', named=True)
	return render(request, 'app/profession_skills.html', {'year': year, 'skills': year_skills, 'prof': year.profession})


def last_vacancies(request):
	global cache, last_fetch
	if last_fetch > timezone.now() - datetime.timedelta(minutes=5):
		return render(request, 'app/last_vacancies.html', {'content': cache})
	response = requests.get(HH_ENDPOINT + '?text=Верстальщик&per_page=10', headers=ua)
	if response.status_code != 200:
		return HttpResponseBadRequest()
	last_fetch = timezone.now()
	ids = [item.get('id') for item in response.json()['items']]
	cache = []
	for _id in ids:
		resp = requests.get(HH_ENDPOINT + f'{_id}/', headers=ua)
		if resp.status_code == 200:
			json = resp.json()
			sort = int(datetime.datetime.strptime(json["published_at"][:-5].replace('T', " "), '%Y-%m-%d %H:%M:%S').timestamp())
			json['sort'] = sort
			cache.append(json)
	cache = list(sorted(cache, key=lambda x: int(x['sort']), reverse=True))
	return render(request, 'app/last_vacancies.html', {'content': cache})


def prepare_rows(prof):
	media_folder = MEDIA_ROOT.split('/')[-1]
	title = ['Название профессии: ', prof.name]
	descs = ProfessionDescriptionBlock.objects.filter(profession=prof)
	description = ['Описания', ]
	if descs.count() == 0:
		description.append('Нет описаний')
	else:
		clean = re.compile('<.*?>')
		for desc in descs:
			desc_title = desc.title
			desc_description = re.sub(clean, '', desc.description)
			desc_description = desc_description.replace('&nbsp;', ' ').replace('&mdash;', '--')
			description.append(f'{desc_title}:\n{desc_description}')
	graphs = ['Графики:', ]
	graphs_rows = []
	graph1 = YearDiagram.objects.filter(profession=prof).first()
	graph2 = VacanciesDiagram.objects.filter(profession=prof).first()
	graph3 = CitiesDiagram.objects.filter(profession=prof).first()
	graph4 = CitiesVacanciesDiagram.objects.filter(profession=prof).first()
	annotate = {
		0: 'year_diagram',
		1: 'vacancies_diagram',
		2: 'cities_diagram',
		3: 'cities_vacancies_diagram',
	}
	for i, diag in enumerate([graph1, graph2, graph3, graph4]):
		if diag:
			folder = f'{annotate[i]}/{prof.id}.jpg'
			save_folder = media_folder + '/' + folder
			link = f'{DOMAIN}{MEDIA_URL}{folder}'
			try:
				with open(save_folder, 'wb') as file:
					file.write(diag.image)
			except:
				link = f'Нет данных'
			graphs_rows.append([diag.diagram_title, link])
		else:
			graphs_rows.append([f'График {i + 1}: ', 'Не загружен'])
	skills = ['Навыки:', ]
	skills_rows = []
	years = Year.objects.filter(profession=prof)
	if years.count() == 0:
		skills_rows = [['Нет данных', ], ]
	else:
		for year in years:
			year_str = str(year.year) + " год: "
			skills_db = Skill.objects.filter(year=year).order_by('-weight')
			if skills_db.count() == 0:
				skills_rows.append([year_str, 'Нет данных'])
			else:
				skills_str = ', '.join([f'{skl.name} - {skl.weight}%' for skl in skills_db])
				skills_rows.append([year_str, skills_str])
	return [title, description, ['', ], graphs, *graphs_rows, ['', ], skills, *skills_rows]


def generate_csv(dest, rows):
	new_file = xlsxwriter.Workbook(dest)
	sheet = new_file.add_worksheet('Информация')
	for r, row in enumerate(rows):
		for c, item in enumerate(row):
			sheet.write(r, c, item)
	new_file.close()


def generate_pdf(dest, rows):
	doc_path = dest[:-4] + '.doc'
	doc = docx.Document()
	for r, row in enumerate(rows):
		if r == 0:
			head = row[0] + row[1]
			doc.add_heading(head, 0)
		elif row == ['', ]:
			doc.add_page_break()
		else:

			head = row.pop(0)
			doc.add_heading(head, 1)
			for in_row in row:
				if DOMAIN in in_row:
					link = in_row.replace(DOMAIN, '')[1:]
					picture = doc.add_picture(link, width=Inches(5))

				elif '%' in in_row:
					items = in_row.split('%, ')
					items[-1] = items[-1][:-1]
					table = doc.add_table(rows=1, cols=2)
					header = table.rows[0].cells
					header[0].text = 'Навык'
					header[1].text = 'Востребованность'
					for item in items:
						i1, i2 = item.split(' - ')
						row_cells = table.add_row().cells
						row_cells[0].text = i1
						row_cells[1].text = i2 + '%'
				else:
					doc.add_paragraph(in_row)
	doc.save(doc_path)
	comtypes.CoInitialize()
	word = comtypes.client.CreateObject('Word.Application')
	doc = word.Documents.Open(str(BASE_DIR / doc_path))
	doc.SaveAs(str(BASE_DIR / dest), FileFormat=wdFormatPDF)
	doc.Close()
	word.Quit()


def download_csv(request, prof_id):
	if not (prof := Profession.objects.filter(id=prof_id, is_visible=True).first()):
		return HttpResponseNotFound()

	media_folder = MEDIA_ROOT.split('\\')[-1]
	file_folder = f'{media_folder}/cache/{prof.id}.xls'
	rows = prepare_rows(prof)
	generate_csv(file_folder, rows)

	return FileResponse(open(file_folder, 'rb'))


def download_pdf(request, prof_id):
	if not (prof := Profession.objects.filter(id=prof_id, is_visible=True).first()):
		return HttpResponseNotFound()

	media_folder = MEDIA_ROOT.split('\\')[-1]
	file_folder = f'{media_folder}/cache/{prof.id}.pdf'
	rows = prepare_rows(prof)

	generate_pdf(file_folder, rows)
	return FileResponse(open(file_folder, 'rb'))